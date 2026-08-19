from converters.base import BaseConverter
import PyPDF2

class DocxConverter(BaseConverter):
    name = 'Word Document (DOCX)'
    output_format = 'docx'
    
    def convert(self, input_path: str, output_path: str, options: dict = None) -> bool:
        try:
            from docx import Document
            from docx.shared import Inches
            
            options = options or {}
            
            doc = Document()
            
            with open(input_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                
                for i, page in enumerate(reader.pages):
                    text = page.extract_text()
                    
                    if i > 0:
                        doc.add_page_break()
                    
                    doc.add_heading(f'Page {i+1}', level=2)
                    
                    for line in text.split('\n'):
                        if line.strip():
                            doc.add_paragraph(line.strip())
            
            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Error converting to DOCX: {e}")
            return False
